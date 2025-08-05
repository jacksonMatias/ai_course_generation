"""
Markdown → DOCX aplicando los estilos de style_map.py y tu
plantilla Word (course_template.docx).

Autor : AI-Course-Generator
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

import markdown as md
from bs4 import BeautifulSoup, Comment
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.templates.style_map import STYLE_MAP

LOGGER = logging.getLogger(__name__)


class MarkdownToWordConverter:
    """Convierte Markdown en DOCX con estilos corporativos."""

    # ------------------------------------------------------------- #
    # API público                                                   #
    # ------------------------------------------------------------- #
    def convert(
        self,
        markdown_text: str,
        output_path: str | Path,
        *,
        template_path: Optional[str | Path] = None,
    ) -> Path:
        output_path = Path(output_path)

        # 1. Documento base (con plantilla si se proporciona)
        doc = Document(str(template_path)) if template_path else Document()

        # 2. Si usamos plantilla, limpiar el cuerpo pero mantener estilos
        if template_path:
            self._wipe_template_body(doc)

        # 3. Fuente/colores por defecto (Normal)
        self._set_base_font(
            doc,
            name="Noto Sans",
            size_pt=10,
            rgb="2C2C30",
        )

        # 4. Markdown → HTML
        html = md.markdown(
            markdown_text,
            extensions=["tables", "fenced_code", "codehilite"],
        )

        # 5. HTML → Word
        soup = BeautifulSoup(html, "html.parser")
        pending_override: str | None = None

        for el in soup.recursiveChildGenerator():
            # Capturar comentarios de estilo
            if isinstance(el, Comment) and el.strip().lower().startswith("style:"):
                pending_override = el.split(":", 1)[1].strip()
                continue

            # Ignorar strings sueltos
            if not getattr(el, "name", None):
                continue

            tag = el.name.lower()

            # ============ Encabezados ==============
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = int(tag[1])
                para = doc.add_paragraph(el.get_text(strip=True))
                self._apply_style_safe(para, STYLE_MAP.get(tag, f"Heading {level}"))

            # ============ Párrafos =================
            elif tag == "p":
                para = doc.add_paragraph(el.get_text())
                if pending_override:
                    self._apply_style_safe(para, STYLE_MAP.get(pending_override, pending_override))
                    pending_override = None
                else:
                    self._apply_style_safe(para, STYLE_MAP.get("p", "Normal"))

            # ============ Listas ===================
            elif tag in {"ul", "ol"}:
                style = STYLE_MAP.get(tag, "List Bullet")
                for li in el.find_all("li", recursive=False):
                    li_para = doc.add_paragraph(li.get_text(strip=True))
                    self._apply_style_safe(li_para, style)

            # ============ Bloque de código =========
            elif tag == "pre":
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(el.get_text())
                code_run.font.name = "Consolas"
                code_run.font.size = Pt(9)

            # ============ Tablas ===================
            elif tag == "table":
                self._convert_table(el, doc)

        doc.save(output_path)
        LOGGER.info("DOCX saved → %s", output_path)
        return output_path

    # ------------------------------------------------------------- #
    # Helpers                                                       #
    # ------------------------------------------------------------- #
    @staticmethod
    def _wipe_template_body(doc: Document) -> None:
        """
        Elimina todo el contenido del cuerpo de la plantilla pero
        mantiene los estilos, headers y footers.
        """
        try:
            body = doc.element.body

            # Guardar sectPr (propiedades de sección) si existe
            sectPr = body.find(qn("w:sectPr"))

            # Eliminar todos los elementos del cuerpo
            for child in list(body):
                if child.tag != qn("w:sectPr"):  # No eliminar propiedades de sección
                    body.remove(child)

            # Añadir un párrafo vacío para que el documento no esté completamente vacío
            empty_p = OxmlElement("w:p")
            body.insert(0, empty_p)

            # Restaurar sectPr al final si existía
            if sectPr is not None:
                body.append(sectPr)

        except Exception as e:
            LOGGER.warning("No se pudo limpiar el cuerpo de la plantilla: %s", e)

    @staticmethod
    def _apply_style_safe(paragraph, style_name: str) -> None:
        """
        Aplica un estilo de forma segura. Si el estilo no existe,
        usa 'Normal' como fallback.
        """
        try:
            paragraph.style = style_name
        except KeyError:
            LOGGER.warning("Estilo '%s' no encontrado, usando 'Normal'", style_name)
            paragraph.style = "Normal"

    @staticmethod
    def _set_base_font(doc: Document, *, name: str, size_pt: int, rgb: str) -> None:
        try:
            font = doc.styles["Normal"].font
            font.name = name
            font.size = Pt(size_pt)
            font.color.rgb = RGBColor.from_string(rgb)
        except Exception as e:
            LOGGER.warning("No se pudo configurar la fuente base: %s", e)

    @staticmethod
    def _convert_table(table_el, doc: Document) -> None:
        rows = table_el.find_all("tr", recursive=False)
        if not rows:
            return

        cols = rows[0].find_all(["td", "th"], recursive=False)
        table = doc.add_table(rows=len(rows), cols=len(cols))

        try:
            table.style = "Table Grid"
        except KeyError:
            LOGGER.warning("Estilo 'Table Grid' no encontrado")

        for r_idx, row in enumerate(rows):
            cells = row.find_all(["td", "th"], recursive=False)
            for c_idx, cell in enumerate(cells):
                if c_idx < len(cols):  # Evitar índices fuera de rango
                    table.cell(r_idx, c_idx).text = cell.get_text(strip=True)
